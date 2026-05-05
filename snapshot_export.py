"""
Collect wizard UI inputs into a structured snapshot for persistence, export (DOCX/PDF/JSON),
and future RAG / LLM context.
"""

from __future__ import annotations

import io
import json
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Mapping

from funnel_finance import FunnelFinanceState
from xml.sax.saxutils import escape

DATA_DIR = Path(__file__).resolve().parent / "data"
SCHEMA = "wisespending.wizard_snapshot.v1"

# Own · Step 1 median check — must match `app.py` radio labels exactly.
OWN_MEDIAN_CONFIRM_YES = "Yes — matches my market expectations"
OWN_MEDIAN_CONFIRM_UNSURE = "Unsure — still continue to the mortgage model"

# Match app.py car spend ceiling (15% take-home) used in the wizard UI — not injected into the pie when lines are blank.
CAR_SPEND_CEILING_FRAC = 0.15


def disk_cache_enabled() -> bool:
    """When false (WISESPENDING_DISABLE_DISK_CACHE=1), skip JSON files — required for multi-user public hosting."""
    v = os.environ.get("WISESPENDING_DISABLE_DISK_CACHE", "").strip().lower()
    return v not in ("1", "true", "yes", "on")


def child_monthly_from_line_items(ss: Mapping[str, Any]) -> float:
    """Sum of child wizard line items — single source for allocation when ``dash_child_monthly`` might be stale."""
    tui = float(ss.get("dash_child_tuition_monthly", 0) or 0)
    care = float(ss.get("dash_child_childcare_monthly", 0) or 0)
    ins = float(ss.get("dash_child_insurance_monthly", 0) or 0)
    act = float(ss.get("dash_child_activities_monthly", 0) or 0)
    ent_c = float(ss.get("dash_child_entertainment_monthly", 0) or 0)
    clo = float(ss.get("dash_child_clothing_monthly", 0) or 0)
    oth = float(ss.get("dash_child_other_monthly", 0) or 0)
    return tui + care + ins + act + ent_c + clo + oth


def entertainment_monthly_from_line_items(ss: Mapping[str, Any]) -> float:
    """Sum of entertainment line items — same idea as ``child_monthly_from_line_items``."""
    out_e = float(ss.get("dash_ent_out_monthly", 0) or 0)
    med = float(ss.get("dash_ent_media_monthly", 0) or 0)
    trp = float(ss.get("dash_ent_trips_monthly", 0) or 0)
    ply = float(ss.get("dash_ent_play_monthly", 0) or 0)
    soc = float(ss.get("dash_ent_social_monthly", 0) or 0)
    ot_e = float(ss.get("dash_ent_other_monthly", 0) or 0)
    return out_e + med + trp + ply + soc + ot_e


def owner_all_in_monthly_typed(ss: Any) -> float:
    """
    Owner **actual all-in** monthly (``dash_mortgage``).

    Streamlit may omit ``dash_mortgage`` from session when the Housing step does not mount
    that ``number_input`` (e.g. **Own** is binding but the UI is on the **Rent** tab).
    The key can also exist with a default **0** before the user edits — then prefer
    ``dash_mortgage_cached`` when it is still positive (see ``_sync_dash_mortgage_cache`` in ``app.py``).
    """
    cached = 0.0
    try:
        cached = float(ss.get("dash_mortgage_cached", 0) or 0)
    except (TypeError, AttributeError):
        cached = 0.0
    live: float | None = None
    try:
        if "dash_mortgage" in ss:
            live = float(ss.get("dash_mortgage") or 0)
    except (TypeError, AttributeError):
        live = None
    if live is not None and live > 0:
        return live
    if cached > 0:
        return cached
    if live is not None:
        return live
    return 0.0


def refresh_session_derived_totals(ss: Any) -> None:
    """Roll up child / entertainment / car line items into monthly totals (mutates mutable session-like dict)."""
    if not (hasattr(ss, "__setitem__") and hasattr(ss, "get")):
        return

    ss["dash_child_monthly"] = int(round(child_monthly_from_line_items(ss)))

    ss["dash_ent"] = int(round(entertainment_monthly_from_line_items(ss)))

    status = ss.get("dash_car_status", "No car")
    if status != "No car":
        pay = (
            float(ss.get("dash_car_payment_monthly", 0) or 0)
            if status in ("Lease", "Finance")
            else 0.0
        )
        dep = (
            float(ss.get("dash_car_depreciation_monthly", 0) or 0)
            if status == "Own outright"
            else 0.0
        )
        ins_m = float(ss.get("dash_car_insurance_monthly", 0) or 0)
        fuel_m = float(ss.get("dash_car_fuel_monthly", 0) or 0)
        maint_m = float(ss.get("dash_car_maintenance_monthly", 0) or 0)
        ss["dash_car_monthly"] = int(round(pay + dep + ins_m + fuel_m + maint_m))
    else:
        ss["dash_car_monthly"] = 0


def effective_housing_monthly(ss: Mapping[str, Any]) -> float:
    """
    Housing for cashflow / pie: typed rent or owned all-in wins first.

    Own Step 1 rules when those are still zero:

    - **Yes — matches my market expectations:** use ZIP **owner-cost median (adjusted)** (`b_o`).
    - **Unsure** then **Verify → mortgage model:** use the **modeled baseline** until you fill
      **Your actual monthly all-in** (`dash_mortgage`); once filled, that amount wins (handled above).

    Otherwise fall back to the Rent/Own worksheet side and benchmarks as before.
    """
    rent = float(ss.get("dash_rent", 0) or 0)
    mort = owner_all_in_monthly_typed(ss)
    modeled = float(ss.get("dash_housing_model_own_monthly", 0) or 0)
    own2 = bool(ss.get("dash_own_mortgage_step", False))
    b_r = float(ss.get("dash_housing_benchmark_rent_mo", 0) or 0)
    b_o = float(ss.get("dash_housing_benchmark_owner_mo", 0) or 0)
    side = str(ss.get("dash_housing_worksheet_side", "rent") or "rent").lower()
    if side not in ("rent", "own"):
        side = "rent"

    confirm = str(ss.get("dash_own_median_confirm") or "")
    from_unsure = bool(ss.get("dash_own_came_from_unsure", False))

    if rent > 0 and mort > 0:
        return max(rent, mort)
    if rent > 0:
        return rent
    if mort > 0:
        return mort

    if from_unsure and own2:
        if modeled > 0:
            return modeled
        if b_o > 0:
            return b_o
        if b_r > 0:
            return b_r
        return 0.0

    if confirm == OWN_MEDIAN_CONFIRM_YES and b_o > 0:
        return b_o

    if side == "own":
        if own2 and modeled > 0:
            return modeled
        if b_o > 0:
            return b_o
        if b_r > 0:
            return b_r
        return 0.0

    if b_r > 0:
        return b_r
    if own2 and modeled > 0:
        return modeled
    if b_o > 0:
        return b_o
    return 0.0


def housing_worksheet_rent_subview_monthly(ss: Mapping[str, Any]) -> float:
    """
    Housing $ for the **Rent** tab banner: **your** rent+utilities when typed, else the **ZIP gross-rent benchmark**
    the tab is showing (same as session ``dash_housing_benchmark_rent_mo`` when filled from ACS).
    """
    rent = float(ss.get("dash_rent", 0) or 0)
    if rent > 0:
        return rent
    return max(0.0, float(ss.get("dash_housing_benchmark_rent_mo", 0) or 0))


def housing_worksheet_own_subview_monthly(ss: Mapping[str, Any]) -> float:
    """
    Housing $ for the **Own** tab banner — only the own-side story (ignores typed rent for this display):

    - Typed **owned all-in** (``dash_mortgage``) wins when **> 0**.
    - **Unsure** + mortgage step: **modeled** baseline until all-in is filled, then benchmarks as fallbacks.
    - **Yes** on ZIP medians: **ZIP owner-cost median** when all-in still **$0**.
    - Else on mortgage step: **modeled** if available, else owner/rent benchmarks.
    """
    mort = owner_all_in_monthly_typed(ss)
    if mort > 0:
        return mort
    modeled = float(ss.get("dash_housing_model_own_monthly", 0) or 0)
    own2 = bool(ss.get("dash_own_mortgage_step", False))
    b_o = float(ss.get("dash_housing_benchmark_owner_mo", 0) or 0)
    b_r = float(ss.get("dash_housing_benchmark_rent_mo", 0) or 0)
    confirm = str(ss.get("dash_own_median_confirm") or "")
    from_unsure = bool(ss.get("dash_own_came_from_unsure", False))

    if from_unsure and own2:
        if modeled > 0:
            return modeled
        if b_o > 0:
            return b_o
        if b_r > 0:
            return b_r
        return 0.0

    if confirm == OWN_MEDIAN_CONFIRM_YES and b_o > 0:
        return b_o

    if own2 and modeled > 0:
        return modeled
    if b_o > 0:
        return b_o
    if b_r > 0:
        return b_r
    return 0.0


def housing_expense_for_math_monthly(ss: Mapping[str, Any]) -> float:
    """
    Housing line **for JSON math + allocation** — follows the **Rent vs Own sub-page the user chose**
    on the Housing step (``dash_housing_committed_subtab``, synced from ``dash_housing_ui_subtab`` whenever
    that module renders). Falls back to ``dash_housing_worksheet_side`` if unset.

    - **rent** → ``housing_worksheet_rent_subview_monthly`` (typed rent + utilities, else ZIP benchmark).
    - **own** → ``dash_mortgage`` when **> 0**; else modeled baseline if **> 0**; else ``housing_worksheet_own_subview_monthly``.

    This can differ from ``effective_housing_monthly`` (e.g. that helper may take max(rent, mort) when both are typed).
    """
    sub = str(ss.get("dash_housing_committed_subtab") or ss.get("dash_housing_worksheet_side", "rent") or "rent").lower()
    if sub not in ("rent", "own"):
        sub = "rent"
    if sub == "rent":
        return float(housing_worksheet_rent_subview_monthly(ss))
    mort = owner_all_in_monthly_typed(ss)
    if mort > 0:
        return mort
    modeled = float(ss.get("dash_housing_model_own_monthly", 0) or 0)
    if modeled > 0:
        return modeled
    return float(housing_worksheet_own_subview_monthly(ss))


def housing_expense_math_rule_applied(ss: Mapping[str, Any]) -> str:
    """Short machine-readable tag describing which branch ``housing_expense_for_math_monthly`` used."""
    sub = str(ss.get("dash_housing_committed_subtab") or ss.get("dash_housing_worksheet_side", "rent") or "rent").lower()
    if sub not in ("rent", "own"):
        sub = "rent"
    if sub == "rent":
        return "committed_rent_worksheet_subview"
    mort = owner_all_in_monthly_typed(ss)
    if mort > 0:
        return "committed_own_actual_all_in_monthly"
    modeled = float(ss.get("dash_housing_model_own_monthly", 0) or 0)
    if modeled > 0:
        return "committed_own_modeled_when_all_in_zero"
    return "committed_own_worksheet_subview_fallback"


def car_expense_math_rule_applied(ss: Mapping[str, Any]) -> str:
    """Tag for mobility line — same sources as ``effective_transport_monthly``."""
    if str(ss.get("dash_car_status", "No car") or "No car") == "No car":
        return "public_transit_shared_mobility_no_car"
    return "vehicle_total_estimated_monthly"


def worksheet_transport_monthly_usd(ss: Mapping[str, Any]) -> float:
    """
    Mobility dollars **after** ``refresh_session_derived_totals`` has updated ``dash_car_monthly``.

    Does **not** call ``refresh`` — use inside ``living_five_amounts_usd_monthly`` so Housing and Car
    lines are read from one refresh pass (avoids subtle reorder bugs when Child lines change).
    """
    status = ss.get("dash_car_status", "No car")
    if status == "No car":
        return max(0.0, float(ss.get("dash_transit_monthly", 0) or 0))
    return max(0.0, float(ss.get("dash_car_monthly", 0) or 0))


def _vehicle_line_items_sum_usd(ss: Mapping[str, Any]) -> float:
    """Same rollup as ``refresh_session_derived_totals`` uses for ``dash_car_monthly`` (read-only)."""
    status = str(ss.get("dash_car_status", "No car") or "No car")
    if status == "No car":
        return 0.0
    pay = (
        float(ss.get("dash_car_payment_monthly", 0) or 0)
        if status in ("Lease", "Finance")
        else 0.0
    )
    dep = (
        float(ss.get("dash_car_depreciation_monthly", 0) or 0)
        if status == "Own outright"
        else 0.0
    )
    ins_m = float(ss.get("dash_car_insurance_monthly", 0) or 0)
    fuel_m = float(ss.get("dash_car_fuel_monthly", 0) or 0)
    maint_m = float(ss.get("dash_car_maintenance_monthly", 0) or 0)
    return max(0.0, pay + dep + ins_m + fuel_m + maint_m)


def _transport_living_usd_monthly(ss: Mapping[str, Any]) -> float:
    """
    Car / transit line for the living-five bundle — align with worksheet **and** raw line items.

    When the dashboard (or chat) renders without the Car step mounted, ``dash_car_monthly`` can lag
    one tick behind typed vehicle lines; taking the max fixes pie / ledger vs tracker.
    """
    t_ws = float(worksheet_transport_monthly_usd(ss))
    status = str(ss.get("dash_car_status", "No car") or "No car")
    if status == "No car":
        t_in = max(0.0, float(ss.get("dash_transit_monthly", 0) or 0))
        return max(t_ws, t_in)
    rolled = _vehicle_line_items_sum_usd(ss)
    dc = max(0.0, float(ss.get("dash_car_monthly", 0) or 0))
    return max(t_ws, rolled, dc)


def _child_living_usd_monthly(ss: Mapping[str, Any]) -> float:
    """Prefer line-item sum; fall back to ``dash_child_monthly`` when line widgets are not mounted."""
    c_li = float(child_monthly_from_line_items(ss))
    if c_li > 1e-9:
        return c_li
    return max(0.0, float(ss.get("dash_child_monthly", 0) or 0))


def _grocery_living_usd_monthly(ss: Mapping[str, Any]) -> float:
    """``dash_grocery`` and grocery tracker pin should match; use max so the pie never drops grocery $."""
    g = max(0.0, float(ss.get("dash_grocery", 0) or 0))
    t = max(0.0, float(ss.get("_wf_grocery_tracker_spend_usd", 0) or 0))
    return max(g, t)


def _entertainment_living_usd_monthly(ss: Mapping[str, Any]) -> float:
    """Prefer six line items; fall back to tracker / ``dash_ent`` when Fun step is not mounted."""
    e_li = float(entertainment_monthly_from_line_items(ss))
    if e_li > 1e-9:
        return e_li
    e_trk = max(0.0, float(ss.get("_wf_entertainment_tracker_spend_usd", 0) or 0))
    e_tot = max(0.0, float(ss.get("dash_ent", 0) or 0))
    return max(e_trk, e_tot)


def effective_transport_monthly(ss: Mapping[str, Any]) -> float:
    """
    Mobility for worksheet, pie, and export — one number, no placeholders:

    - **No car:** ``dash_transit_monthly`` (Public transit & shared mobility).
    - **Lease / Finance / Own outright:** ``dash_car_monthly`` = sum of the vehicle lines
      (same as **Total estimated monthly vehicle cost** in the UI).

    If every relevant line is **0**, the result is **0**. If any line is **> 0**, the rolled-up
    total (transit alone, or sum of vehicle lines) is what we record.
    """
    if hasattr(ss, "__setitem__") and hasattr(ss, "get"):
        refresh_session_derived_totals(ss)
    return worksheet_transport_monthly_usd(ss)


def session_take_home_usd_monthly(ss: Any) -> float:
    """Read ``monthly_take_home`` from Streamlit session or a plain mapping (robust to .get vs attribute)."""
    v: Any = None
    try:
        v = ss.get("monthly_take_home", None)  # type: ignore[union-attr]
    except Exception:
        v = None
    if v is None:
        try:
            v = ss["monthly_take_home"]  # type: ignore[index]
        except Exception:
            v = getattr(ss, "monthly_take_home", 0)
    try:
        return float(v or 0)
    except (TypeError, ValueError):
        return 0.0


def living_five_amounts_usd_monthly(ss: Mapping[str, Any]) -> tuple[float, float, float, float, float]:
    """
    First five waterfall lines — same bundle as ``compute_takehome_allocation``.

    Car / child / grocery / fun use **fallbacks** so the pie and ledger stay aligned with wizard trackers
    when this runs from the **dashboard** (widgets for those steps may not be on the page that rerun).
    """
    refresh_session_derived_totals(ss)
    return (
        float(housing_expense_for_math_monthly(ss)),
        float(_transport_living_usd_monthly(ss)),
        float(_child_living_usd_monthly(ss)),
        float(_grocery_living_usd_monthly(ss)),
        float(_entertainment_living_usd_monthly(ss)),
    )


def _housing_pin_matches_saved(h_saved: float, h_current: float) -> bool:
    """Housing at pin time vs current math — allow a few dollars / basis-point drift (ZIP vs rounding)."""
    a, b = float(h_saved), float(h_current)
    return abs(a - b) <= max(5.0, 0.001 * max(abs(a), abs(b), 1.0))


def _housing_pin_matches_for_wizard_replay(h_saved: float, h_current: float) -> bool:
    """
    When re-applying ``_wf_*_meta`` from a prior wizard step, housing can drift **more** than the tight
    math-line tolerance (user switched Rent/Own, ZIP benchmark refresh, or committed sub-tab changed) while
    **income + Car/Child/Grocery/Fun** in the pin are still the right bundle. If we require ``_housing_pin_matches_saved``
    only, Car/Child stay **$0** on the pie and dollars wrongly pile into **Final savings**.
    """
    a, b = float(h_saved), float(h_current)
    return abs(a - b) <= max(200.0, 0.05 * max(abs(a), abs(b), 1.0))


def capture_wizard_step_spending_at_next(ss: Any, slug: str) -> None:
    """
    When the user clicks **Next →** on a dashboard wizard step, snapshot that step’s monthly spend into
    ``_wf_next_pie_*`` so the pie / ledger can read stable USD even when this step’s widgets are not mounted.

    ``slug`` is ``housing`` | ``car`` | ``child`` | ``grocery`` | ``entertainment`` (same as tracker slugs).
    """
    if not hasattr(ss, "__setitem__"):
        return
    refresh_session_derived_totals(ss)
    inc = session_take_home_usd_monthly(ss)
    ss["_wf_next_pie_income_snapshot"] = float(inc)
    if slug == "housing":
        ss["_wf_next_pie_housing_usd"] = float(housing_expense_for_math_monthly(ss))
    elif slug == "car":
        ss["_wf_next_pie_car_usd"] = float(_transport_living_usd_monthly(ss))
    elif slug == "child":
        ss["_wf_next_pie_child_usd"] = float(_child_living_usd_monthly(ss))
    elif slug == "grocery":
        ss["_wf_next_pie_grocery_usd"] = float(_grocery_living_usd_monthly(ss))
    elif slug == "entertainment":
        ss["_wf_next_pie_fun_usd"] = float(_entertainment_living_usd_monthly(ss))


def apply_next_pie_cache_max_to_living_five(
    ss: Mapping[str, Any],
    income: float,
    housing: float,
    transport: float,
    child: float,
    grocery: float,
    ent: float,
) -> tuple[float, float, float, float, float]:
    """
    Merge explicit **Next →** snapshots (``_wf_next_pie_*``) into the five living lines using ``max`` so
    worksheet re-entries can still increase a line but never drop below what the user committed when leaving a step.
    """
    snap = ss.get("_wf_next_pie_income_snapshot")
    if snap is None or not _income_pin_matches_saved(float(snap), float(income)):
        return (float(housing), float(transport), float(child), float(grocery), float(ent))
    h = float(housing)
    t = float(transport)
    c = float(child)
    g = float(grocery)
    e = float(ent)
    if ss.get("_wf_next_pie_housing_usd") is not None:
        h = max(h, float(ss["_wf_next_pie_housing_usd"]))
    if ss.get("_wf_next_pie_car_usd") is not None:
        t = max(t, float(ss["_wf_next_pie_car_usd"]))
    if ss.get("_wf_next_pie_child_usd") is not None:
        c = max(c, float(ss["_wf_next_pie_child_usd"]))
    if ss.get("_wf_next_pie_grocery_usd") is not None:
        g = max(g, float(ss["_wf_next_pie_grocery_usd"]))
    if ss.get("_wf_next_pie_fun_usd") is not None:
        e = max(e, float(ss["_wf_next_pie_fun_usd"]))
    return (h, t, c, g, e)


def _income_pin_matches_saved(inc_saved: float, inc_current: float) -> bool:
    """
    Take-home at pin time vs session ``monthly_take_home``.

    Pins must not require **2¢** equality: Streamlit ``number_input`` / JSON reload often differ from the
    floats stored in ``_wf_*_meta`` by **dollars**, which used to skip the Fun/Child merge while the
    entertainment **S** pin still matched — Car/Child stayed **$0** on the pie and dollars hid in **Final savings**.
    """
    a, b = float(inc_saved), float(inc_current)
    return abs(a - b) <= max(2.0, 0.001 * max(abs(a), abs(b), 1.0))


def _meta_float_prefix(meta: Any, n: int) -> list[float] | None:
    """Read first ``n`` floats from a JSON-restored list or session tuple."""
    if meta is None or not isinstance(meta, (list, tuple)) or len(meta) < n:
        return None
    try:
        return [float(meta[i]) for i in range(n)]
    except (TypeError, ValueError, IndexError):
        return None


def savings_before_investment_usd_monthly(
    ss: Mapping[str, Any],
    *,
    _living_five: tuple[float, float, float, float, float] | None = None,
) -> float:
    """
    Monthly **S** (savings before investment): ``take-home − (five living lines)`` after applying the same
    **Next →** row cache (``_wf_next_pie_*``) as ``compute_takehome_allocation``.
    """
    if _living_five is None:
        housing, transport, child, grocery, ent = living_five_amounts_usd_monthly(ss)
    else:
        housing, transport, child, grocery, ent = _living_five
    income = session_take_home_usd_monthly(ss)
    housing, transport, child, grocery, ent = apply_next_pie_cache_max_to_living_five(
        ss, float(income), float(housing), float(transport), float(child), float(grocery), float(ent)
    )
    return float(income) - float(housing) - float(transport) - float(child) - float(grocery) - float(ent)


def compute_takehome_allocation(
    ss: Mapping[str, Any],
    *,
    _living_five: tuple[float, float, float, float, float] | None = None,
) -> dict[str, float]:
    """
    Monthly take-home split (USD) — **waterfall**:

    1. **Living:** housing (``housing_expense_for_math_monthly`` — Rent vs **Own sub-page choice** ``dash_housing_committed_subtab``) →
       car/transit (``effective_transport_monthly``) → child → groceries → entertainment.
    2. **Savings before investment (S):** ``S = take-home − living`` after Fun/Child/Car pin hydration and
       ``_wf_next_pie_*`` **Next →** snapshots (same five lines as the pie).
    3. **Emergency savings (this month):** ``ef`` — dollars you move this month (pie row).
    4. **Long-term investing (hypothetical %):** ``max(0, S − ef) × (slider % / 100)`` — slider applies to **cash after**
       this month’s emergency move from the post-living pool, not to full take-home.
    5. **Final savings / unallocated:** ``(S − ef) − invest`` = take-home − living − ef − invest — reconciles the pie.
       **401(k) match** stays reflection-only (no dollar row).

    Pass ``_living_five`` when you already called ``living_five_amounts_usd_monthly`` in the same render
    (avoids a duplicate ``refresh_session_derived_totals``).
    """
    if _living_five is None:
        housing, transport, child, grocery, ent = living_five_amounts_usd_monthly(ss)
    else:
        housing, transport, child, grocery, ent = _living_five
    income = session_take_home_usd_monthly(ss)

    # Re-hydrate Car/Child/Grocery/Fun from Fun-step meta when valid, else Child/Car pins; then merge **Next →**
    # snapshots (``_wf_next_pie_*``) so the pie never drops dollars the user committed when leaving a step.
    meta_e = ss.get("_wf_entertainment_pool_meta")
    merged_from_fun_pin = False
    seq6 = _meta_float_prefix(meta_e, 6)
    if seq6 is not None:
        inc_m, h_m, c_m, ch_m, g_m, fun_m = seq6
        if _income_pin_matches_saved(inc_m, float(income)) and _housing_pin_matches_for_wizard_replay(
            h_m, float(housing)
        ):
            transport, child, grocery, ent = c_m, ch_m, g_m, fun_m
            merged_from_fun_pin = True
    if not merged_from_fun_pin:
        # No usable Fun meta (e.g. new browser session before Fun) — recover **Car + Child** from Child-step pin.
        seq4 = _meta_float_prefix(ss.get("_wf_pool_after_child_meta"), 4)
        if seq4 is not None:
            ic, hc, tc, kc = seq4
            if _income_pin_matches_saved(ic, float(income)) and _housing_pin_matches_for_wizard_replay(
                hc, float(housing)
            ):
                if float(transport) < 1e-6 and tc > 1e-6:
                    transport = tc
                if float(child) < 1e-6 and kc > 1e-6:
                    child = kc
        # Car step only wrote ``_wf_car_pool_meta`` until Child — recover **Car** if still zero after Child pin.
        if float(transport) < 1e-6:
            seq3 = _meta_float_prefix(ss.get("_wf_car_pool_meta"), 3)
            if seq3 is not None:
                ic2, hc2, car_c = seq3
                if (
                    _income_pin_matches_saved(ic2, float(income))
                    and _housing_pin_matches_for_wizard_replay(hc2, float(housing))
                    and car_c > 1e-6
                ):
                    transport = car_c

    housing, transport, child, grocery, ent = apply_next_pie_cache_max_to_living_five(
        ss, float(income), float(housing), float(transport), float(child), float(grocery), float(ent)
    )
    savings_before_investment = float(income) - float(housing) - float(transport) - float(child) - float(
        grocery
    ) - float(ent)

    ef_save = max(0.0, float(ss.get("dash_ef_monthly_save", 0) or 0))
    cash_after_emergency = savings_before_investment - ef_save

    inv_pct = int(ss.get("dash_invest_pct", 0) or 0)
    inv_pct = max(0, min(40, inv_pct))
    invest_base = max(0.0, cash_after_emergency)
    invest_mo = invest_base * inv_pct / 100.0

    final = cash_after_emergency - invest_mo

    out: dict[str, float] = {
        "Housing (effective — rent, owned, model, or ZIP benchmark)": housing,
        "Car or transit (effective — lines or guideline)": transport,
        "Child-related": child,
        "Groceries (food at home)": grocery,
        "Entertainment & discretionary": ent,
        "Emergency savings (this month)": ef_save,
        "Long-term investing (hypothetical %)": invest_mo,
        "Final savings / unallocated": final,
    }
    drift = float(income) - sum(float(v) for v in out.values())
    if abs(drift) > 1e-6:
        out["Final savings / unallocated"] = float(out["Final savings / unallocated"]) + drift
    return out


TAKEHOME_ALLOCATION_PIE_KEYS: tuple[str, ...] = (
    "Housing (effective — rent, owned, model, or ZIP benchmark)",
    "Car or transit (effective — lines or guideline)",
    "Child-related",
    "Groceries (food at home)",
    "Entertainment & discretionary",
    "Emergency savings (this month)",
    "Long-term investing (hypothetical %)",
    "Final savings / unallocated",
)

# Stable IDs + display labels + colors — same order as ``TAKEHOME_ALLOCATION_PIE_KEYS`` (sidebar: Housing → … → Invest bucket rows).
WATERFALL_PIE_LINE_ID: tuple[str, ...] = (
    "housing",
    "car_transit",
    "child",
    "grocery",
    "entertainment",
    "emergency_savings_month",
    "long_term_invest_hypothetical",
    "final_savings_unallocated",
)
WATERFALL_PIE_LABEL_SHORT: tuple[str, ...] = (
    "Housing",
    "Car / transit",
    "Child",
    "Groceries",
    "Fun",
    "Emergency save",
    "Invest (%)",
    "Final savings",
)

# Monochrome emerald palette (ledger + exports align with in-app pie).
WATERFALL_PIE_COLOR_HEX: tuple[str, ...] = (
    "#064e3b",
    "#065f46",
    "#047857",
    "#059669",
    "#10b981",
    "#34d399",
    "#6ee7b7",
    "#a7f3d0",
)

WATERFALL_PIE_LEDGER_SCHEMA = "wisespending.waterfall_pie_ledger.v1"
WATERFALL_PIE_LEDGER_FILENAME = "waterfall_pie_ledger.json"

WATERFALL_RAG_CONTEXT_SCHEMA = "wisespending.waterfall_rag_context.v1"
WATERFALL_RAG_CONTEXT_FILENAME = "waterfall_rag_context.json"


def _housing_ledger_provenance(ss: Mapping[str, Any]) -> tuple[str, str, str]:
    rule = housing_expense_math_rule_applied(ss)
    sub = str(ss.get("dash_housing_committed_subtab") or ss.get("dash_housing_worksheet_side", "rent") or "rent").lower()
    if sub not in ("rent", "own"):
        sub = "rent"
    if sub == "rent":
        rent = float(ss.get("dash_rent", 0) or 0)
        if rent > 0:
            return rule, "user_input", "Rent & utilities (typed)."
        br = float(ss.get("dash_housing_benchmark_rent_mo", 0) or 0)
        if br > 0:
            return rule, "system_estimate", "ZIP-area gross-rent benchmark (ACS-style estimate)."
        return rule, "none", "No typed rent and no benchmark yet."
    mort = owner_all_in_monthly_typed(ss)
    if mort > 0:
        return rule, "user_input", "Owned home: actual all-in monthly (typed)."
    modeled = float(ss.get("dash_housing_model_own_monthly", 0) or 0)
    if modeled > 0:
        return rule, "system_estimate", "Owned home: modeled baseline until actual all-in is filled."
    b_o = float(ss.get("dash_housing_benchmark_owner_mo", 0) or 0)
    if b_o > 0:
        return rule, "system_estimate", "ZIP-area owner-cost median or worksheet fallback."
    return rule, "none", "Own path with no modeled, benchmark, or typed all-in yet."


def _car_ledger_provenance(ss: Mapping[str, Any]) -> tuple[str, str, str]:
    rule = car_expense_math_rule_applied(ss)
    if rule == "public_transit_shared_mobility_no_car":
        t = float(ss.get("dash_transit_monthly", 0) or 0)
        if t > 0:
            return rule, "user_input", "No car — public transit & shared mobility (typed)."
        return rule, "none", "No car — mobility lines at $0."
    return rule, "user_input", "Vehicle — sum of payment/depreciation/insurance/fuel/maintenance (wizard lines)."


def _waterfall_line_provenance(ss: Mapping[str, Any], row_index: int) -> tuple[str, str, str]:
    """(rule_or_reason_tag, amount_source, short_notes) for one waterfall / pie row."""
    if row_index == 0:
        return _housing_ledger_provenance(ss)
    if row_index == 1:
        return _car_ledger_provenance(ss)
    if row_index == 2:
        return "child_line_items", "user_input", "Child costs — sum of wizard line items."
    if row_index == 3:
        g = float(ss.get("dash_grocery", 0) or 0)
        if g > 0:
            return "dash_grocery", "user_input", "Groceries (food at home) — user estimate."
        return "dash_grocery", "none", "Groceries not entered ($0)."
    if row_index == 4:
        return "entertainment_line_items", "user_input", "Fun — sum of six entertainment line items."
    if row_index == 5:
        ef = float(ss.get("dash_ef_monthly_save", 0) or 0)
        if ef > 0:
            return "dash_ef_monthly_save", "user_input", "Emergency savings — this month’s planned move."
        return "dash_ef_monthly_save", "none", "No emergency move entered ($0)."
    if row_index == 6:
        return "invest_slider_on_post_emergency", "derived", "Long-term invest (hypothetical): slider % × max($0, S − emergency)."
    return "remainder_plus_drift", "derived", "Final cushion: take-home minus all rows above; absorbs rounding drift for the pie."


def build_waterfall_pie_ledger(ss: Mapping[str, Any]) -> dict[str, Any]:
    """
    Canonical JSON-ready record of **every dollar row** in the take-home waterfall / pie.

    Amounts are exactly what ``compute_takehome_allocation`` applies (including drift on the last row).
    ``amount_source`` is **user_input** (typed/slider), **system_estimate** (benchmarks/model), **derived**
    (formula from prior rows), or **none** ($0 / not applicable).
    """
    refresh_session_derived_totals(ss)
    alloc = compute_takehome_allocation(ss)
    income = session_take_home_usd_monthly(ss)
    keys = TAKEHOME_ALLOCATION_PIE_KEYS
    raw = [float(alloc[k]) for k in keys]
    drift = float(income) - sum(raw)
    drift_applied = 0.0
    if abs(drift) > 1e-6 and income > 0:
        raw[-1] += drift
        drift_applied = float(drift)
    lines_out: list[dict[str, Any]] = []
    for i, k in enumerate(keys):
        rule, src, notes = _waterfall_line_provenance(ss, i)
        usd = round(raw[i], 2)
        lines_out.append(
            {
                "waterfall_index": i,
                "id": WATERFALL_PIE_LINE_ID[i],
                "pie_row_key": k,
                "label_short": WATERFALL_PIE_LABEL_SHORT[i],
                "usd_monthly": usd,
                "amount_source": src,
                "rule_or_reason": rule,
                "notes": notes,
                "color_hex": WATERFALL_PIE_COLOR_HEX[i],
            }
        )
    return {
        "schema": WATERFALL_PIE_LEDGER_SCHEMA,
        "updated_at_utc": datetime.now(timezone.utc).isoformat(),
        "monthly_take_home_usd": round(float(income), 2),
        "drift_applied_to_final_usd": round(drift_applied, 4),
        "lines": lines_out,
        "lines_sum_usd_monthly": round(sum(raw), 2),
    }


def persist_waterfall_pie_ledger(ledger: dict[str, Any]) -> Path | None:
    """Write ``data/waterfall_pie_ledger.json`` — single source used to render the spending pie."""
    if not disk_cache_enabled():
        return None
    try:
        DATA_DIR.mkdir(parents=True, exist_ok=True)
        path = waterfall_pie_ledger_path()
        path.write_text(json.dumps(ledger, indent=2, ensure_ascii=False), encoding="utf-8")
        return path
    except OSError:
        return None


def waterfall_pie_ledger_path() -> Path:
    """On-disk ledger consumed when ``prefer_ledger_file`` is used in ``build_waterfall_for_pie``."""
    return DATA_DIR / WATERFALL_PIE_LEDGER_FILENAME


def load_waterfall_pie_ledger_from_disk() -> dict[str, Any] | None:
    """
    Read ``data/waterfall_pie_ledger.json`` via ``Path.read_text`` + ``json.loads``.

    Returns ``None`` if the file is missing, invalid JSON, wrong schema, or the wrong number of lines.
    """
    path = waterfall_pie_ledger_path()
    if not path.is_file():
        return None
    try:
        raw: Any = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, UnicodeError, json.JSONDecodeError):
        return None
    if not isinstance(raw, dict):
        return None
    if raw.get("schema") != WATERFALL_PIE_LEDGER_SCHEMA:
        return None
    lines = raw.get("lines")
    if not isinstance(lines, list) or len(lines) != len(TAKEHOME_ALLOCATION_PIE_KEYS):
        return None
    for row in lines:
        if not isinstance(row, dict):
            return None
        for key in ("id", "label_short", "usd_monthly", "color_hex"):
            if key not in row:
                return None
    try:
        float(raw["monthly_take_home_usd"])
    except (KeyError, TypeError, ValueError):
        return None
    return raw


def takehome_amounts_from_ledger_dict(ledger: Mapping[str, Any]) -> tuple[float, list[float]]:
    """
    Take-home and the eight waterfall USD amounts for display — same drift rule as the pie/table
    (last line absorbs ``monthly_take_home_usd − sum(lines)`` when rounding leaves a gap).

    ``ledger`` may be the on-disk ``waterfall_pie_ledger.json`` payload or any dict with the same
    ``monthly_take_home_usd`` + ``lines`` shape.
    """
    inc = float(ledger["monthly_take_home_usd"])
    lines = ledger["lines"]
    vals = [float(x["usd_monthly"]) for x in lines]  # type: ignore[index]
    drift = inc - sum(vals)
    if abs(drift) > 1e-6 and inc > 0:
        vals[-1] += drift
    return inc, vals


def waterfall_for_pie_from_ledger_dict(ledger: dict[str, Any]) -> dict[str, Any]:
    """Shape returned by ``build_waterfall_for_pie`` (pie + table), built only from a ledger dict."""
    lines: list[dict[str, Any]] = []
    for entry in ledger["lines"]:
        row = entry  # validated dict in ``load_waterfall_pie_ledger_from_disk``
        lines.append(
            {
                "id": str(row["id"]),
                "pie_row_key": str(row.get("pie_row_key", "")),
                "label_short": str(row["label_short"]),
                "usd_monthly": float(row["usd_monthly"]),
                "color_hex": str(row["color_hex"]),
            }
        )
    inc = float(ledger["monthly_take_home_usd"])
    sm = ledger.get("lines_sum_usd_monthly")
    if sm is None:
        sm = sum(float(x["usd_monthly"]) for x in lines)
    return {
        "monthly_take_home_usd": inc,
        "lines": lines,
        "lines_sum_usd_monthly": float(sm),
        "ledger_schema": ledger.get("schema"),
        "ledger_updated_at_utc": ledger.get("updated_at_utc"),
        "pie_data_source": "ledger_file",
    }


def build_waterfall_for_pie(ss: Mapping[str, Any], *, prefer_ledger_file: bool = False) -> dict[str, Any]:
    """
    Pie + UI document from ``build_waterfall_pie_ledger``, persisted to JSON.

    If ``prefer_ledger_file`` is ``True``, load ``data/waterfall_pie_ledger.json`` first (external edits).
    When the file is missing or invalid, falls back to a live ledger build and still persists when allowed.
    """
    if prefer_ledger_file:
        disk_ledger = load_waterfall_pie_ledger_from_disk()
        if disk_ledger is not None:
            return waterfall_for_pie_from_ledger_dict(disk_ledger)
    ledger = build_waterfall_pie_ledger(ss)
    try:
        persist_waterfall_pie_ledger(ledger)
    except OSError:
        pass
    lines: list[dict[str, Any]] = []
    for entry in ledger["lines"]:
        lines.append(
            {
                "id": entry["id"],
                "pie_row_key": entry["pie_row_key"],
                "label_short": entry["label_short"],
                "usd_monthly": float(entry["usd_monthly"]),
                "color_hex": entry["color_hex"],
            }
        )
    return {
        "monthly_take_home_usd": float(ledger["monthly_take_home_usd"]),
        "lines": lines,
        "lines_sum_usd_monthly": float(ledger["lines_sum_usd_monthly"]),
        "ledger_schema": ledger.get("schema"),
        "ledger_updated_at_utc": ledger.get("updated_at_utc"),
        "pie_data_source": "live_session",
    }


# Wizard slug → which waterfall row(s) (0-based, same order as ``TAKEHOME_ALLOCATION_PIE_KEYS``) the
# in-app **Take-home tracker** highlights for that page. Invest and chat steps have no tracker (pie / chat UI instead).
WIZARD_SLUG_TO_WATERFALL_TRACKER_ROWS: dict[str, tuple[int, ...]] = {
    "housing": (0,),
    "car": (1,),
    "child": (2,),
    "grocery": (3,),
    "entertainment": (4,),
}


def compute_wizard_surplus(ss: Mapping[str, Any]) -> tuple[float, float, float, dict[str, float]]:
    """Final cushion after living, emergency move, and hypothetical long-term invest (% of max(0, savings-before-investment − emergency))."""
    alloc = compute_takehome_allocation(ss)
    income = session_take_home_usd_monthly(ss)
    final_savings = alloc["Final savings / unallocated"]
    total_allocated = income - final_savings
    breakdown = {k: round(float(v), 2) for k, v in alloc.items()}
    return final_savings, income, round(total_allocated, 2), breakdown


def build_funnel_finance_state_from_wizard_session(ss: Mapping[str, Any]) -> tuple[FunnelFinanceState, list[str]]:
    """
    Strict funnel (no emergency in these two formulas): map questionnaire + worksheets
    into ``FunnelFinanceState``.

    ``investing`` is ``(slider % / 100) × max(0, S)`` with ``S`` = income minus the five
    living lines — parallel to “% of surplus before invest” without mixing in emergency.
    """
    housing, car, child, grocery, entertainment = living_five_amounts_usd_monthly(ss)
    income = session_take_home_usd_monthly(ss)
    inv_pct = int(ss.get("dash_invest_pct", 0) or 0)
    inv_pct = max(0, min(40, inv_pct))
    lf = (housing, car, child, grocery, entertainment)
    s_before = savings_before_investment_usd_monthly(ss, _living_five=lf)
    investing_requested = max(0.0, s_before) * (inv_pct / 100.0)
    st = FunnelFinanceState(
        income=income,
        housing=housing,
        car=car,
        child=child,
        grocery=grocery,
        entertainment=entertainment,
        investing=investing_requested,
    )
    st.saving_before_investing = float(s_before)
    _lang = str(ss.get("ui_lang", "en") or "en")
    if _lang not in ("en", "es", "zh"):
        _lang = "en"
    warnings = st.warnings_for_surplus_and_invest_cap(_lang)
    return st, warnings


def collect_wizard_snapshot(ss: Mapping[str, Any]) -> dict[str, Any]:
    surplus, inc, out, brk = compute_wizard_surplus(ss)
    ff_st, ff_ws = build_funnel_finance_state_from_wizard_session(ss)
    ff_dict = {k: round(float(v), 2) for k, v in ff_st.as_dict().items()}
    prefer_ledger_pie = bool(ss.get("dash_pie_prefer_ledger_file", False))
    waterfall_for_pie = build_waterfall_for_pie(ss, prefer_ledger_file=prefer_ledger_pie)
    keys = list(TAKEHOME_ALLOCATION_PIE_KEYS)
    living_sum = sum(float(brk[keys[i]]) for i in range(5))
    lf5 = tuple(float(brk[keys[i]]) for i in range(5))
    savings_before_investment = round(savings_before_investment_usd_monthly(ss, _living_five=lf5), 2)
    housing_math = round(float(brk[keys[0]]), 2)
    car_math = round(float(brk[keys[1]]), 2)
    child_m = round(float(brk[keys[2]]), 2)
    grocery_m = round(float(brk[keys[3]]), 2)
    ent_m = round(float(brk[keys[4]]), 2)
    saving_before_investment_math = savings_before_investment
    inv_pct_i = int(ss.get("dash_invest_pct", 0) or 0)
    inv_pct_i = max(0, min(40, inv_pct_i))
    post_emergency = round(savings_before_investment - float(brk[keys[5]]), 2)
    invest_base_pos = max(0.0, float(post_emergency))
    invest_pct_base = round(invest_base_pos, 2)
    savings_after_long_term_invest_fraction = round(invest_base_pos * (1.0 - inv_pct_i / 100.0), 2)
    return {
        "schema": SCHEMA,
        "updated_at_utc": datetime.now(timezone.utc).isoformat(),
        "questionnaire": {
            "monthly_take_home": inc,
            "take_home_income_usd_monthly": round(float(inc), 2),
            "family": str(ss.get("family", "")),
            "spending_vibe": str(ss.get("spending_vibe", "")),
        },
        "cashflow_math_inputs_usd_monthly": {
            "take_home_income": round(float(inc), 2),
            "housing_expense": housing_math,
            "housing_rule_applied": housing_expense_math_rule_applied(ss),
            "car_or_transit_expense": car_math,
            "car_or_transit_rule_applied": car_expense_math_rule_applied(ss),
            "child_expense": child_m,
            "grocery_expense": grocery_m,
            "entertainment_expense": ent_m,
            "saving_before_investment": saving_before_investment_math,
        },
        "computed": {
            "monthly_surplus_per_wizard_buckets": round(surplus, 2),
            "tracked_monthly_spending_sum": round(out, 2),
            "spending_breakdown": {k: round(v, 2) for k, v in brk.items()},
            "waterfall_living_expenses_sum_monthly": round(living_sum, 2),
            "waterfall_savings_before_investment_monthly": savings_before_investment,
            "waterfall_savings_after_long_term_invest_fraction_usd_monthly": savings_after_long_term_invest_fraction,
            "waterfall_take_home_minus_living_monthly": savings_before_investment,
            "waterfall_cash_after_emergency_save_monthly": post_emergency,
            "waterfall_invest_slider_percent_base_usd_monthly": invest_pct_base,
            "waterfall_for_pie": waterfall_for_pie,
            "waterfall_pie_ledger": {
                "schema": WATERFALL_PIE_LEDGER_SCHEMA,
                "disk_relative_path": f"data/{WATERFALL_PIE_LEDGER_FILENAME}",
                "updated_at_utc": waterfall_for_pie.get("ledger_updated_at_utc"),
                "description": "Full line-level ledger (user vs system vs derived) is written next to this snapshot; pie slices are built from that file.",
            },
            "funnel_finance_strict_usd_monthly": ff_dict,
            "funnel_finance_strict_warnings": ff_ws,
        },
        "housing": {
            "zip_input": str(ss.get("dash_housing_zip", "") or ""),
            "rent_monthly": float(ss.get("dash_rent", 0) or 0),
            "own_on_mortgage_step": bool(ss.get("dash_own_mortgage_step", False)),
            "own_baseline_mode": ss.get("dash_own_baseline_mode"),
            "own_median_confirmation": ss.get("dash_own_median_confirm"),
            "own_home_price": float(ss.get("dash_own_home_price", 0) or 0),
            "own_down_payment_pct": float(ss.get("dash_own_down_pct", 0) or 0),
            "own_interest_rate_pct": float(ss.get("dash_own_rate_pct", 0) or 0),
            "own_loan_term_years": int(ss.get("dash_own_term", 30) or 30),
            "own_property_tax_style": ss.get("dash_own_property_kind"),
            "own_actual_all_in_housing_monthly": round(owner_all_in_monthly_typed(ss), 2),
            "acs_benchmark_rent_monthly": float(ss.get("dash_housing_benchmark_rent_mo", 0) or 0),
            "acs_benchmark_owner_cost_monthly": float(ss.get("dash_housing_benchmark_owner_mo", 0) or 0),
            "modeled_owner_baseline_monthly": float(ss.get("dash_housing_model_own_monthly", 0) or 0),
            "worksheet_estimate_side_rent_or_own": ss.get("dash_housing_worksheet_side"),
            "housing_ui_subtab_rent_or_own": ss.get("dash_housing_ui_subtab"),
            "housing_committed_subtab_rent_or_own": ss.get("dash_housing_committed_subtab"),
            "effective_worksheet_housing_monthly": round(effective_housing_monthly(ss), 2),
            "housing_expense_for_cashflow_math_usd_monthly": housing_math,
            "housing_expense_for_cashflow_math_rule": housing_expense_math_rule_applied(ss),
        },
        "transportation": {
            "cars_band": ss.get("dash_car_household_cars"),
            "car_status": ss.get("dash_car_status"),
            "mobility_worksheet_basis": (
                "public_transit_monthly"
                if (ss.get("dash_car_status") or "No car") == "No car"
                else "vehicle_lines_total_monthly"
            ),
            "effective_worksheet_transport_monthly": round(effective_transport_monthly(ss), 2),
            "car_or_transit_expense_for_cashflow_math_usd_monthly": car_math,
            "car_or_transit_expense_for_cashflow_math_rule": car_expense_math_rule_applied(ss),
            "transit_monthly_if_no_car": float(ss.get("dash_transit_monthly", 0) or 0),
            "lease_or_finance_payment_monthly": float(ss.get("dash_car_payment_monthly", 0) or 0),
            "own_depreciation_estimate_monthly": float(ss.get("dash_car_depreciation_monthly", 0) or 0),
            "insurance_monthly": float(ss.get("dash_car_insurance_monthly", 0) or 0),
            "fuel_monthly": float(ss.get("dash_car_fuel_monthly", 0) or 0),
            "maintenance_monthly": float(ss.get("dash_car_maintenance_monthly", 0) or 0),
            "vehicle_total_monthly": float(ss.get("dash_car_monthly", 0) or 0),
            "depreciation_illustration_vehicle_price": float(ss.get("dash_car_dep_example_price", 0) or 0),
        },
        "child": {
            "tuition_monthly": float(ss.get("dash_child_tuition_monthly", 0) or 0),
            "childcare_monthly": float(ss.get("dash_child_childcare_monthly", 0) or 0),
            "insurance_monthly": float(ss.get("dash_child_insurance_monthly", 0) or 0),
            "activities_monthly": float(ss.get("dash_child_activities_monthly", 0) or 0),
            "entertainment_monthly": float(ss.get("dash_child_entertainment_monthly", 0) or 0),
            "clothing_essentials_monthly": float(ss.get("dash_child_clothing_monthly", 0) or 0),
            "other_monthly": float(ss.get("dash_child_other_monthly", 0) or 0),
            "child_total_monthly": float(ss.get("dash_child_monthly", 0) or 0),
        },
        "grocery": {"at_home_monthly": float(ss.get("dash_grocery", 0) or 0)},
        "entertainment": {
            "eating_drinks_out_monthly": float(ss.get("dash_ent_out_monthly", 0) or 0),
            "subscriptions_media_monthly": float(ss.get("dash_ent_media_monthly", 0) or 0),
            "events_travel_monthly": float(ss.get("dash_ent_trips_monthly", 0) or 0),
            "hobbies_shopping_monthly": float(ss.get("dash_ent_play_monthly", 0) or 0),
            "social_family_monthly": float(ss.get("dash_ent_social_monthly", 0) or 0),
            "other_monthly": float(ss.get("dash_ent_other_monthly", 0) or 0),
            "entertainment_total_monthly": float(ss.get("dash_ent", 0) or 0),
        },
        "invest_reflections": {
            "emergency_fund_target_months_of_take_home": int(ss.get("dash_ef_target_months", 6) or 6),
            "emergency_fund_current_balance_usd": float(ss.get("dash_ef_current_balance", 0) or 0),
            "emergency_fund_planned_monthly_contribution_usd": float(ss.get("dash_ef_monthly_save", 0) or 0),
            "employer_401k_match_self_report": ss.get("dash_401k_match_status"),
            "hypothetical_long_term_invest_pct_slider_0_40": int(ss.get("dash_invest_pct", 10) or 10),
            "hypothetical_long_term_invest_usd_monthly": round(float(brk[keys[6]]), 2),
        },
    }


def persist_snapshot(snap: dict[str, Any]) -> Path | None:
    if not disk_cache_enabled():
        return None
    try:
        DATA_DIR.mkdir(parents=True, exist_ok=True)
        path = DATA_DIR / "wizard_snapshot.json"
        path.write_text(json.dumps(snap, indent=2, ensure_ascii=False), encoding="utf-8")
        return path
    except OSError:
        return None


def snapshot_to_rag_plaintext(
    snap: dict[str, Any],
    *,
    waterfall_ledger: dict[str, Any] | None = None,
) -> str:
    """Plaintext bundle for demo chat + exports: authoritative waterfall ledger first, then full snapshot."""
    parts: list[str] = [
        "# Wise spending — RAG / assistant context\n",
        "Use **WATERFALL_PIE_LEDGER** first: it is the exact monthly USD that feed the pie/waterfall "
        "(housing, car/transit, child, groceries, fun, emergency save, hypothetical invest %, final savings). "
        "Each line includes `amount_source`, `rule_or_reason`, and `notes`.\n",
        "The **wizard_snapshot** block repeats worksheets and computed fields; figures are user estimates and may be incomplete.\n\n",
    ]
    if waterfall_ledger:
        parts.append("## WATERFALL_PIE_LEDGER (canonical monthly rows)\n\n")
        parts.append(json.dumps(waterfall_ledger, indent=2, ensure_ascii=False))
        parts.append("\n\n## WIZARD_SNAPSHOT (full worksheets + computed)\n\n")
    else:
        parts.append("## WIZARD_SNAPSHOT (full worksheets + computed)\n\n")
    parts.append(json.dumps(snap, indent=2, ensure_ascii=False))
    return "".join(parts)


def build_docx_bytes(snap: dict[str, Any]) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Wise spending — Wizard inputs", level=0)
    p = doc.add_paragraph()
    run = p.add_run(f"Exported (UTC): {snap.get('updated_at_utc', '')}")
    run.italic = True
    doc.add_paragraph(f"Schema: {snap.get('schema', '')}")

    def add_section(title: str, blob: Any) -> None:
        doc.add_heading(title, level=1)
        if isinstance(blob, dict):
            for k, v in blob.items():
                doc.add_paragraph(f"{k}: {v}", style="List Bullet")
        else:
            doc.add_paragraph(str(blob))

    for key in (
        "questionnaire",
        "cashflow_math_inputs_usd_monthly",
        "computed",
        "housing",
        "transportation",
        "child",
        "grocery",
        "entertainment",
        "invest_reflections",
    ):
        if key in snap:
            add_section(key.replace("_", " ").title(), snap[key])

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_pdf_bytes(snap: dict[str, Any]) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph(escape("Wise spending — Wizard inputs"), styles["Title"]))
    story.append(Paragraph(escape(str(snap.get("updated_at_utc", ""))), styles["Italic"]))
    story.append(Spacer(1, 12))

    def add_blob(title: str, obj: Any) -> None:
        story.append(Paragraph(escape(title), styles["Heading2"]))
        if isinstance(obj, dict):
            for k, v in obj.items():
                story.append(Paragraph(escape(f"{k}: {v}"), styles["Normal"]))
        else:
            story.append(Paragraph(escape(str(obj)), styles["Normal"]))
        story.append(Spacer(1, 8))

    for key in (
        "questionnaire",
        "cashflow_math_inputs_usd_monthly",
        "computed",
        "housing",
        "transportation",
        "child",
        "grocery",
        "entertainment",
        "invest_reflections",
    ):
        if key in snap:
            add_blob(key, snap[key])

    doc.build(story)
    return buf.getvalue()


# --- Disk cache: survives browser tab close / Streamlit rerun (same machine) -----------

USER_INPUTS_SCHEMA = "wisespending.wizard_user_inputs.v1"
USER_INPUTS_FILENAME = "wizard_user_inputs.json"


WIZARD_USER_INPUT_KEYS: tuple[str, ...] = (
    "onboarding_complete",
    "onboarding_step",
    "dashboard_step",
    "monthly_take_home",
    "family",
    "spending_vibe",
    "dash_housing_zip",
    "dash_housing_benchmark_rent_mo",
    "dash_housing_benchmark_owner_mo",
    "dash_housing_model_own_monthly",
    "dash_housing_worksheet_side",
    "dash_housing_ui_subtab",
    "dash_housing_committed_subtab",
    "dash_rent",
    "dash_mortgage",
    "dash_mortgage_cached",
    "dash_own_mortgage_step",
    "dash_own_baseline_mode",
    "dash_own_median_confirm",
    "dash_own_home_price",
    "dash_own_down_pct",
    "dash_own_rate_pct",
    "dash_own_term",
    "dash_own_property_kind",
    "dash_own_came_from_unsure",
    "dash_car_household_cars",
    "dash_car_status",
    "dash_transit_monthly",
    "dash_car_payment_monthly",
    "dash_car_depreciation_monthly",
    "dash_car_insurance_monthly",
    "dash_car_fuel_monthly",
    "dash_car_maintenance_monthly",
    "dash_car_monthly",
    "dash_car_dep_example_price",
    "dash_child_tuition_monthly",
    "dash_child_childcare_monthly",
    "dash_child_insurance_monthly",
    "dash_child_activities_monthly",
    "dash_child_entertainment_monthly",
    "dash_child_clothing_monthly",
    "dash_child_other_monthly",
    "dash_child_monthly",
    "dash_grocery",
    "dash_ent_out_monthly",
    "dash_ent_media_monthly",
    "dash_ent_trips_monthly",
    "dash_ent_play_monthly",
    "dash_ent_social_monthly",
    "dash_ent_other_monthly",
    "dash_ent",
    "dash_invest_pct",
    "dash_ef_target_months",
    "dash_ef_current_balance",
    "dash_ef_monthly_save",
    "dash_401k_match_status",
    "dash_pie_prefer_ledger_file",
    "advisor_messages",
    "llm_model",
    "llm_base_url",
    # Waterfall pins (JSON stores tuples as lists — hydrate back to tuple in ``app._apply_disk_user_input_cache``).
    "_wf_car_pool_meta",
    "_wf_car_pool_usd",
    "_wf_pool_after_child_meta",
    "_wf_pool_after_child_usd",
    "_wf_grocery_pool_meta",
    "_wf_grocery_pool_usd",
    "_wf_grocery_tracker_spend_usd",
    "_wf_entertainment_pool_meta",
    "_wf_entertainment_pool_usd",
    "_wf_entertainment_tracker_spend_usd",
    # Per-step **Next →** snapshots for the pie (same USD as leaving that module).
    "_wf_next_pie_income_snapshot",
    "_wf_next_pie_housing_usd",
    "_wf_next_pie_car_usd",
    "_wf_next_pie_child_usd",
    "_wf_next_pie_grocery_usd",
    "_wf_next_pie_fun_usd",
)


def user_inputs_cache_path() -> Path:
    return DATA_DIR / USER_INPUTS_FILENAME


def persist_user_inputs_to_disk(ss: Mapping[str, Any]) -> Path | None:
    """Write whitelisted session keys to JSON (UTF-8) for reload after tab/browser close."""
    if not disk_cache_enabled():
        return None
    vals: dict[str, Any] = {}
    for k in WIZARD_USER_INPUT_KEYS:
        if k not in ss:
            continue
        v = ss[k]
        try:
            json.dumps(v)
        except (TypeError, ValueError):
            continue
        vals[k] = v
    payload = {
        "schema": USER_INPUTS_SCHEMA,
        "updated_at_utc": datetime.now(timezone.utc).isoformat(),
        "values": vals,
    }
    try:
        DATA_DIR.mkdir(parents=True, exist_ok=True)
        p = user_inputs_cache_path()
        p.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
        return p
    except OSError:
        return None


def load_user_inputs_from_disk() -> dict[str, Any] | None:
    if not disk_cache_enabled():
        return None
    p = user_inputs_cache_path()
    if not p.is_file():
        return None
    try:
        raw = json.loads(p.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError, UnicodeDecodeError):
        return None
    if raw.get("schema") != USER_INPUTS_SCHEMA:
        return None
    out = raw.get("values")
    return out if isinstance(out, dict) else None


LLM_ADVISOR_CONTEXT_SCHEMA = "wisespending.llm_advisor_context.v1"
LLM_ADVISOR_CONTEXT_FILENAME = "llm_advisor_context.json"

ADVISOR_LLM_SYSTEM_INSTRUCTION = """
You are the user's personal finance companion: warm, candid, and precise — like a close friend who happens to have strong money habits and a calm planning style. You are **not** a licensed fiduciary; give educational, balanced perspectives and clearly say when a CPA, attorney, or licensed investment advisor should be involved.

Tone: professional but approachable — sound like a knowledgeable advisor who cares; use short paragraphs; no hype or jargon walls. Answer the user's question directly first, then add useful context.

**Chat formatting:** The user only sees this conversation. **USER_FINANCIAL_CONTEXT** is private background for you only. **Never** paste raw JSON, code fences of the payload, or internal field dumps into your reply. Refer to their numbers in plain language (e.g. “your ~$X/mo on housing”).

Grounding: treat **USER_FINANCIAL_CONTEXT** (JSON bundled in this system message) as the user's own worksheet. If it looks incomplete or inconsistent with real life, say so gently and suggest what to double-check.

**Priority fields:** read **`llm_quick_reference`** first — it restates **monthly take-home**, **household family** and **spending vibe** from onboarding, and **`pie_slices_same_order_as_chart`** (USD/mo per slice) matching the on-screen pie chart order.

Then read **`waterfall_pie_ledger`** — authoritative **eight monthly rows** (housing, car/transit, child, groceries, fun, emergency savings this month, long-term invest hypothetical, final savings/unallocated). Each row has **`usd_monthly`**, **`amount_source`** (user_input / system_estimate / derived / none), **`rule_or_reason`**, and **`notes`**. This is the same math as the dashboard pie and is persisted for RAG.

Then use **`wizard_snapshot`** for housing/transport/child/grocery/entertainment **worksheet line items** and **`cashflow_breakdown_usd_monthly`** for named bucket totals.

Cashflow shape in the snapshot: **waterfall** — living (housing through fun) first; **S** = take-home minus that bundle; **emergency** this month comes out of **S**; **hypothetical long-term invest** = slider % × ``max(0, S − emergency)``; **final savings** is what’s left. **401(k) match** in the UI is education/reflection only unless the user models dollars elsewhere.

Typical questions — how to answer:
- **Buying a home / higher housing spend:** translate into monthly impact on cushion, relate to their take-home and existing category totals, mention taxes/insurance/maintenance as common blind spots, and keep arithmetic tied to their numbers.
- **Buying a car / higher transport cost:** compare payment vs total cost of ownership, and tradeoffs against other buckets they already entered.

Always reply in the **same language** the user writes in (e.g. Chinese if they write Chinese).

Do not invent balances or income not present in the context; if something critical is missing, state one clear assumption or ask one focused follow-up.
""".strip()


def build_llm_advisor_context_payload(ss: Mapping[str, Any]) -> dict[str, Any]:
    """Structured JSON + advisor persona text for LLM calls and on-disk handoff."""
    snap = collect_wizard_snapshot(ss)
    _surplus, _inc, _out, brk = compute_wizard_surplus(ss)
    waterfall_ledger = build_waterfall_pie_ledger(ss)
    wf = (snap.get("computed") or {}).get("waterfall_for_pie") or {}
    q = snap.get("questionnaire") or {}
    lines_raw = wf.get("lines") or []
    pie_lines: list[dict[str, Any]] = []
    for x in lines_raw:
        if not isinstance(x, dict):
            continue
        try:
            u = round(float(x.get("usd_monthly", 0)), 2)
        except (TypeError, ValueError):
            u = 0.0
        pie_lines.append(
            {
                "pie_row_key": x.get("pie_row_key"),
                "label_short": x.get("label_short"),
                "usd_monthly": u,
            }
        )
    take_home = wf.get("monthly_take_home_usd")
    if take_home is None:
        take_home = q.get("take_home_income_usd_monthly")
    try:
        take_home_f = round(float(take_home), 2) if take_home is not None else None
    except (TypeError, ValueError):
        take_home_f = None
    llm_quick_reference = {
        "monthly_take_home_usd": take_home_f,
        "household_family": q.get("family"),
        "spending_vibe": q.get("spending_vibe"),
        "pie_data_source": wf.get("pie_data_source"),
        "pie_lines_sum_usd_monthly": wf.get("lines_sum_usd_monthly"),
        "pie_slices_same_order_as_chart": pie_lines,
    }
    return {
        "schema": LLM_ADVISOR_CONTEXT_SCHEMA,
        "updated_at_utc": datetime.now(timezone.utc).isoformat(),
        "system_instruction": ADVISOR_LLM_SYSTEM_INSTRUCTION,
        "llm_quick_reference": llm_quick_reference,
        "waterfall_pie_ledger": waterfall_ledger,
        "wizard_snapshot": snap,
        "cashflow_breakdown_usd_monthly": {k: round(float(v), 2) for k, v in brk.items()},
        "usage_notes": (
            "Bundle `system_instruction` + `llm_quick_reference` + `waterfall_pie_ledger` + `cashflow_breakdown_usd_monthly` + `wizard_snapshot`. "
            "`waterfall_pie_ledger` is the canonical eight-row monthly waterfall (housing through final savings) with provenance per row; "
            "it is also written to `data/waterfall_pie_ledger.json` and bundled in `data/waterfall_rag_context.json` when disk cache is on. "
            "`llm_quick_reference` duplicates onboarding and pie slices for quick scanning. "
            "The snapshot is user-entered; it may omit taxes, medical, debt service, etc. "
            "When `pie_data_source` is `ledger_file`, slices match the saved ledger; otherwise they match the live session (same as the pie UI when the ledger checkbox is off). "
            "Invest slider % applies to `computed.waterfall_invest_slider_percent_base_usd_monthly` (= max(0, S − emergency), see `waterfall_cash_after_emergency_save_monthly`), not to full take-home."
        ),
    }


def llm_advisor_context_path() -> Path:
    return DATA_DIR / LLM_ADVISOR_CONTEXT_FILENAME


def waterfall_rag_context_path() -> Path:
    return DATA_DIR / WATERFALL_RAG_CONTEXT_FILENAME


def build_waterfall_rag_context_document(ss: Mapping[str, Any]) -> dict[str, Any]:
    """Single JSON bundle for RAG: full waterfall ledger + bucket totals (same math as pie)."""
    _surplus, _inc, _out, brk = compute_wizard_surplus(ss)
    ledger = build_waterfall_pie_ledger(ss)
    return {
        "schema": WATERFALL_RAG_CONTEXT_SCHEMA,
        "updated_at_utc": datetime.now(timezone.utc).isoformat(),
        "waterfall_pie_ledger": ledger,
        "cashflow_breakdown_usd_monthly": {k: round(float(v), 2) for k, v in brk.items()},
    }


def persist_waterfall_rag_context(ss: Mapping[str, Any]) -> Path | None:
    """Write ``data/waterfall_rag_context.json`` — ledger + breakdown for chat RAG and external tools."""
    if not disk_cache_enabled():
        return None
    doc = build_waterfall_rag_context_document(ss)
    try:
        DATA_DIR.mkdir(parents=True, exist_ok=True)
        p = waterfall_rag_context_path()
        p.write_text(json.dumps(doc, indent=2, ensure_ascii=False), encoding="utf-8")
        return p
    except OSError:
        return None


def persist_llm_advisor_context(ss: Mapping[str, Any]) -> Path | None:
    """Write `data/llm_advisor_context.json` for chat / external LLM pipelines."""
    if not disk_cache_enabled():
        return None
    payload = build_llm_advisor_context_payload(ss)
    try:
        DATA_DIR.mkdir(parents=True, exist_ok=True)
        p = llm_advisor_context_path()
        p.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
    except OSError:
        return None
    try:
        persist_waterfall_rag_context(ss)
    except OSError:
        pass
    return p


def clear_wizard_disk_caches() -> None:
    if not disk_cache_enabled():
        return
    for name in (
        USER_INPUTS_FILENAME,
        "wizard_snapshot.json",
        LLM_ADVISOR_CONTEXT_FILENAME,
        WATERFALL_PIE_LEDGER_FILENAME,
        WATERFALL_RAG_CONTEXT_FILENAME,
    ):
        p = DATA_DIR / name
        try:
            p.unlink(missing_ok=True)
        except OSError:
            pass
